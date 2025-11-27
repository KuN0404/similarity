import os
import docx
import fitz  # PyMuPDF
import datetime
import re
from django.conf import settings
from nltk.tokenize import sent_tokenize, word_tokenize
from googlesearch import search
from apps.repository.models import RepositoryFile
from apps.plagiarism.models import PlagiarismSettings
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib import colors
from reportlab.lib.units import inch

class PlagiarismService:
    def __init__(self):
        self.threshold = PlagiarismSettings.get_threshold()
        self.matched_sources = []

    def validate_pdf(self, file_path):
        """
        Pre-validation untuk PDF sebelum processing
        Returns: (is_valid, error_message)
        """
        try:
            doc = fitz.open(file_path)
            
            # Check if PDF is empty
            if len(doc) == 0:
                doc.close()
                return (False, "PDF kosong (0 halaman)")
            
            # Quick text check from first page
            first_page = doc[0]
            sample_text = first_page.get_text("text")
            
            doc.close()
            
            # If first page has no text, likely a scan
            if not sample_text or len(sample_text.strip()) < 20:
                return (False, 
                       "PDF kemungkinan berupa scan/gambar. "
                       "Gunakan OCR atau copy-paste teks manual ke form 'Paste Teks'.")
            
            return (True, None)
            
        except fitz.FileDataError:
            return (False, "File PDF corrupt atau tidak valid")
        except fitz.PasswordError:
            return (False, "PDF terproteksi password")
        except Exception as e:
            return (False, f"Error validasi PDF: {str(e)}")

    def extract_text(self, file_path, file_ext):
        """
        Enhanced text extraction with better PDF handling
        """
        text = ""
        try:
            if file_ext == '.pdf':
                # Pre-validate PDF
                is_valid, error_msg = self.validate_pdf(file_path)
                if not is_valid:
                    raise ValueError(error_msg)
                
                text = self._extract_from_pdf(file_path)
            elif file_ext == '.docx':
                text = self._extract_from_docx(file_path)
            
            # Clean and normalize text
            text = self._clean_text(text)
            return text
            
        except Exception as e:
            print(f"Error extracting text: {e}")
            import traceback
            traceback.print_exc()
            raise  # Re-raise to be handled by caller

    def _extract_from_pdf(self, file_path):
        """
        Advanced PDF text extraction - IGNORE IMAGES, EXTRACT TEXT ONLY
        """
        text = ""
        total_text_chars = 0
        total_images = 0
        
        try:
            # Open PDF document
            doc = fitz.open(file_path)
            total_pages = len(doc)
            
            print(f"📄 Processing PDF: {total_pages} pages")
            
            for page_num in range(total_pages):
                page = doc[page_num]
                page_text = ""
                
                try:
                    # Count images on this page (for logging only)
                    image_list = page.get_images()
                    page_images = len(image_list)
                    total_images += page_images
                    
                    # Method 1: Standard text extraction (TEXT ONLY, NO IMAGES)
                    page_text = page.get_text("text")
                    
                    # Method 2: If standard fails, try blocks method (SKIP IMAGE BLOCKS)
                    if not page_text or len(page_text.strip()) < 20:
                        print(f"  Page {page_num+1}: Using blocks method")
                        blocks = page.get_text("blocks")
                        # Filter: block[6] == 0 means TEXT block (not image)
                        text_blocks = [
                            block[4] for block in blocks 
                            if len(block) > 6 and block[6] == 0 and block[4].strip()
                        ]
                        page_text = "\n".join(text_blocks)
                    
                    # Method 3: If still empty, try dict method (TEXT ONLY)
                    if not page_text or len(page_text.strip()) < 20:
                        print(f"  Page {page_num+1}: Using dict method")
                        text_dict = page.get_text("dict")
                        page_text = self._extract_from_dict(text_dict)
                    
                    # Method 4: Try HTML extraction (then strip all tags)
                    if not page_text or len(page_text.strip()) < 20:
                        print(f"  Page {page_num+1}: Using HTML method")
                        html_text = page.get_text("html")
                        # Remove image tags first
                        html_text = re.sub(r'<img[^>]*>', '', html_text, flags=re.IGNORECASE)
                        # Strip all HTML tags
                        page_text = re.sub('<[^<]+?>', '', html_text)
                    
                    # Clean and add page content
                    if page_text and page_text.strip():
                        clean_page_text = page_text.strip()
                        text += clean_page_text + "\n\n"
                        total_text_chars += len(clean_page_text)
                        
                        log_msg = f"  ✓ Page {page_num+1}: {len(clean_page_text)} chars"
                        if page_images > 0:
                            log_msg += f" | {page_images} images (ignored)"
                        print(log_msg)
                    else:
                        if page_images > 0:
                            print(f"  ⚠ Page {page_num+1}: Only images, no text")
                        else:
                            print(f"  ⚠ Page {page_num+1}: No content found")
                
                except Exception as page_error:
                    print(f"  ✗ Page {page_num+1}: Error - {page_error}")
                    continue
                    
            doc.close()
            
            # Final validation
            text = text.strip()
            
            # Summary
            print(f"\n📊 Extraction Summary:")
            print(f"   - Total text extracted: {total_text_chars} characters")
            print(f"   - Total images found: {total_images} (ignored)")
            print(f"   - Pages processed: {total_pages}")
            
            if not text:
                if total_images > 0:
                    raise ValueError(
                        f"PDF berisi {total_images} gambar tetapi tidak ada teks yang dapat diekstrak. "
                        f"Kemungkinan PDF berupa koleksi gambar/scan. "
                        f"Gunakan OCR atau copy-paste teks manual."
                    )
                else:
                    raise ValueError("Tidak ada teks yang dapat diekstrak dari PDF. PDF mungkin kosong.")
            
            if len(text) < 50:
                raise ValueError(
                    f"Teks yang diekstrak terlalu sedikit ({len(text)} karakter). "
                    f"PDF mungkin mostly berisi gambar. "
                    f"Gunakan OCR atau copy-paste teks manual ke form 'Paste Teks'."
                )
            
            print(f"✅ PDF extraction completed: {len(text)} characters (text only)")
            return text
            
        except fitz.FileDataError:
            raise ValueError("File PDF corrupt atau tidak valid. Silakan coba file lain.")
        except fitz.PasswordError:
            raise ValueError("PDF terproteksi password. Silakan hapus password terlebih dahulu.")
        except Exception as e:
            error_msg = str(e)
            if "insufficient text" in error_msg.lower() or "no text" in error_msg.lower() or "gambar" in error_msg.lower():
                raise ValueError(error_msg)
            else:
                raise ValueError(f"Error saat memproses PDF: {error_msg}")

    def _extract_from_dict(self, text_dict):
        """
        Extract text from PyMuPDF dict structure (TEXT ONLY, SKIP IMAGES)
        """
        text = ""
        try:
            if "blocks" in text_dict:
                for block in text_dict["blocks"]:
                    # Skip image blocks (type 1), only process text blocks (type 0)
                    if block.get("type") != 0:
                        continue
                    
                    if "lines" in block:
                        for line in block["lines"]:
                            if "spans" in line:
                                line_text = " ".join([
                                    span.get("text", "") 
                                    for span in line["spans"] 
                                    if span.get("text", "").strip()
                                ])
                                if line_text.strip():
                                    text += line_text.strip() + "\n"
            
            # Remove excessive newlines
            text = re.sub(r'\n\s*\n', '\n\n', text)
            
        except Exception as e:
            print(f"  Error in dict extraction: {e}")
        
        return text

    def _extract_from_docx(self, file_path):
        """
        Extract text from DOCX with paragraph preservation
        """
        try:
            doc = docx.Document(file_path)
            
            # Extract from paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Extract from tables
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            table_texts.append(cell.text.strip())
            
            # Combine all text
            text = '\n'.join(paragraphs)
            if table_texts:
                text += "\n\n" + '\n'.join(table_texts)
            
            if not text or len(text.strip()) < 100:
                raise ValueError("DOCX extraction resulted in insufficient text")
            
            print(f"✓ DOCX extracted successfully: {len(text)} characters")
            return text
            
        except Exception as e:
            print(f"✗ DOCX extraction error: {e}")
            raise ValueError(f"Gagal mengekstrak teks dari DOCX: {str(e)}")

    def _clean_text(self, text):
        """
        Clean and normalize extracted text
        """
        if not text:
            return ""
        
        # Preserve Indonesian characters and common symbols
        allowed_chars = set('abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789 .,!?;:\n\r\t()-"\'/')
        
        # Clean text character by character
        cleaned = []
        for char in text:
            if char in allowed_chars or ord(char) > 127:  # Keep unicode chars (for Indonesian)
                cleaned.append(char)
            else:
                cleaned.append(' ')
        
        text = ''.join(cleaned)
        
        # Normalize whitespace
        text = re.sub(r'[ \t]+', ' ', text)  # Multiple spaces to single
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)  # Max 2 newlines
        
        # Clean each line
        lines = []
        for line in text.split('\n'):
            line = line.strip()
            if line:  # Only keep non-empty lines
                lines.append(line)
        
        text = '\n'.join(lines)
        
        return text.strip()

    def tokenize(self, text):
        """
        Tokenize text into sentences with validation
        """
        if not text or not text.strip():
            return []
        
        try:
            # Use NLTK sentence tokenizer
            sentences = sent_tokenize(text)
            
            # Filter valid sentences
            valid_sentences = []
            for s in sentences:
                s = s.strip()
                # Keep sentences with at least 10 chars and 3 words
                if len(s) > 10 and len(s.split()) >= 3:
                    valid_sentences.append(s)
            
            print(f"✓ Tokenized into {len(valid_sentences)} valid sentences")
            return valid_sentences
            
        except Exception as e:
            print(f"Error tokenizing: {e}")
            # Fallback: split by period
            sentences = text.split('.')
            return [s.strip() + '.' for s in sentences if len(s.strip()) > 10]

    def check_google(self, sentence):
        try:
            query = f'"{sentence}"'
            results = list(search(query, num_results=3, sleep_interval=2))
            return (100, results[0] if results else None) if len(results) > 0 else (0, None)
        except Exception as e:
            print(f"Google search error: {e}")
            return (0, None)

    def check_local(self, sentence):
        best_score = 0
        best_match = None
        indexed_files = RepositoryFile.objects.filter(status='indexed')
        
        if not indexed_files.exists():
            return (0, None)
        
        sentence_lower = sentence.lower()
        sentence_tokens = set(word_tokenize(sentence_lower))
        
        if not sentence_tokens:
            return (0, None)

        for repo_file in indexed_files:
            try:
                txt_path = repo_file.extracted_text_path
                if not txt_path or not os.path.exists(txt_path):
                    continue
                    
                with open(txt_path, 'r', encoding='utf-8') as f:
                    content = f.read().lower()

                if sentence_lower in content:
                    return (100, repo_file)

                content_tokens = set(word_tokenize(content))
                intersection = sentence_tokens.intersection(content_tokens)
                
                if len(sentence_tokens) > 0:
                    overlap_score = (len(intersection) / len(sentence_tokens)) * 100
                else:
                    overlap_score = 0
                
                if overlap_score > best_score:
                    best_score = overlap_score
                    best_match = repo_file
                    
                if best_score == 100:
                    break
                    
            except Exception as e:
                print(f"Error checking local: {e}")
                continue
        
        if best_score >= self.threshold:
            return (best_score, best_match)
        return (0, None)

    def process_check(self, text, source_mode='both'):
        sentences = self.tokenize(text)
        results = []
        local_matches = {}
        internet_matches = set()
        
        total_sentences = len(sentences)
        local_plagiarized = 0
        internet_plagiarized = 0
        
        if total_sentences == 0:
            raise ValueError("Tidak ada kalimat valid yang dapat diperiksa. Pastikan dokumen mengandung teks yang cukup.")
        
        print(f"Checking {total_sentences} sentences with threshold {self.threshold}%")
        
        for idx, sent in enumerate(sentences, 1):
            score_local = 0
            score_internet = 0
            matched_repo = None
            matched_url = None
            
            if source_mode in ['local', 'both']:
                score_local, matched_repo = self.check_local(sent)
                if score_local >= self.threshold:
                    local_plagiarized += 1
                    if matched_repo and matched_repo.id not in local_matches:
                        local_matches[matched_repo.id] = {
                            'id': matched_repo.id,
                            'title': matched_repo.title or 'Unknown',
                            'author': matched_repo.author or 'Unknown',
                            'year': matched_repo.year or 'N/A',
                            'file_path': matched_repo.file.path if matched_repo.file else None,
                            'count': 0
                        }
                    if matched_repo:
                        local_matches[matched_repo.id]['count'] += 1
            
            if source_mode in ['internet', 'both'] and score_local < 100:
                score_internet, matched_url = self.check_google(sent)
                if score_internet >= self.threshold:
                    internet_plagiarized += 1
                    if matched_url:
                        internet_matches.add(matched_url)
            
            final_score = max(score_local, score_internet)
            
            if final_score >= self.threshold:
                result = {
                    'sentence': sent,
                    'score': final_score,
                    'score_local': score_local,
                    'score_internet': score_internet,
                    'source': "Local Repository" if score_local >= score_internet else "Internet",
                }
                
                if matched_repo:
                    result['metadata'] = {
                        'title': matched_repo.title or 'Unknown',
                        'author': matched_repo.author or 'Unknown',
                        'year': matched_repo.year or 'N/A',
                        'repo_id': matched_repo.id
                    }
                elif matched_url:
                    result['metadata'] = {
                        'url': matched_url
                    }
                
                results.append(result)
            
            # Progress indicator
            if idx % 10 == 0:
                print(f"Progress: {idx}/{total_sentences} sentences checked")
        
        similarity_local = int((local_plagiarized / total_sentences) * 100) if total_sentences > 0 else 0
        similarity_internet = int((internet_plagiarized / total_sentences) * 100) if total_sentences > 0 else 0
        
        unique_plagiarized = len(results)
        similarity_global = int((unique_plagiarized / total_sentences) * 100) if total_sentences > 0 else 0
        
        print(f"✓ Check completed: {len(results)} plagiarized sentences found")
        print(f"  Local: {similarity_local}%, Internet: {similarity_internet}%, Global: {similarity_global}%")
        
        return {
            'results': results,
            'similarity_local': similarity_local,
            'similarity_internet': similarity_internet,
            'similarity_global': similarity_global,
            'local_sources': list(local_matches.values()),
            'internet_sources': list(internet_matches)
        }

    def generate_pdf_report(self, original_text, check_results, output_path, filename):
        try:
            results = check_results['results']
            similarity_local = check_results['similarity_local']
            similarity_internet = check_results['similarity_internet']
            similarity_global = check_results['similarity_global']
            local_sources = check_results['local_sources']
            internet_sources = check_results['internet_sources']
            
            doc = SimpleDocTemplate(output_path, pagesize=A4, topMargin=0.5*inch, bottomMargin=0.5*inch)
            story = []
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                textColor=colors.HexColor('#1a1a1a'),
                spaceAfter=12,
                alignment=1,
                fontName='Helvetica-Bold'
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                textColor=colors.HexColor('#333333'),
                spaceAfter=10,
                fontName='Helvetica-Bold'
            )
            
            subheading_style = ParagraphStyle(
                'SubHeading',
                parent=styles['Heading3'],
                fontSize=12,
                textColor=colors.HexColor('#555555'),
                spaceAfter=8,
                fontName='Helvetica-Bold'
            )
            
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=10,
                leading=14
            )
            
            # Title
            story.append(Paragraph("LAPORAN DETEKSI PLAGIARISME", title_style))
            story.append(Spacer(1, 0.2*inch))
            
            # Document info
            story.append(Paragraph(f"<b>Nama File:</b> {filename}", normal_style))
            story.append(Paragraph(f"<b>Tanggal Pemeriksaan:</b> {datetime.datetime.now().strftime('%d-%m-%Y %H:%M')}", normal_style))
            story.append(Paragraph(f"<b>Threshold:</b> {self.threshold}%", normal_style))
            story.append(Spacer(1, 0.3*inch))
            
            # Similarity Index
            story.append(Paragraph("<b>INDEKS SIMILARITAS</b>", heading_style))
            
            similarity_data = [
                ['Kategori', 'Persentase'],
                ['Similaritas Global', f'{similarity_global}%'],
                ['Dari Repository Lokal', f'{similarity_local}%'],
                ['Dari Internet', f'{similarity_internet}%']
            ]
            
            similarity_table = Table(similarity_data, colWidths=[3*inch, 2*inch])
            similarity_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c3e50')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 11),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('FONTNAME', (0, 1), (0, -1), 'Helvetica-Bold'),
            ]))
            story.append(similarity_table)
            story.append(Spacer(1, 0.3*inch))
            
            # Local sources
            if local_sources:
                story.append(Paragraph("<b>SUMBER DARI REPOSITORY LOKAL</b>", heading_style))
                story.append(Spacer(1, 0.1*inch))
                
                source_data = [['No', 'Judul', 'Penulis', 'Tahun', 'Kecocokan']]
                for idx, source in enumerate(local_sources, 1):
                    source_data.append([
                        str(idx),
                        Paragraph(source['title'], normal_style),
                        source['author'],
                        str(source['year']),
                        f"{source['count']} kalimat"
                    ])
                
                source_table = Table(source_data, colWidths=[0.4*inch, 2.5*inch, 1.5*inch, 0.8*inch, 1*inch])
                source_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ]))
                story.append(source_table)
                story.append(Spacer(1, 0.3*inch))
            
            # Internet sources
            if internet_sources:
                story.append(Paragraph("<b>SUMBER DARI INTERNET</b>", heading_style))
                story.append(Spacer(1, 0.1*inch))
                
                for idx, url in enumerate(internet_sources, 1):
                    story.append(Paragraph(f"{idx}. {url}", normal_style))
                    story.append(Spacer(1, 0.05*inch))
                
                story.append(Spacer(1, 0.3*inch))
            
            # Page break
            story.append(PageBreak())
            
            # Detail sentences
            story.append(Paragraph("<b>DETAIL KALIMAT TERDETEKSI</b>", heading_style))
            story.append(Spacer(1, 0.2*inch))
            
            if results:
                detail_data = [['Kalimat', 'Sumber', 'Skor', 'Metadata']]
                for res in results:
                    metadata_text = ""
                    if 'metadata' in res:
                        meta = res['metadata']
                        if 'title' in meta:
                            metadata_text = f"{meta['title']}\n{meta['author']} ({meta['year']})"
                        elif 'url' in meta:
                            metadata_text = meta['url'][:50] + "..."
                    
                    detail_data.append([
                        Paragraph(res['sentence'][:80] + '...' if len(res['sentence']) > 80 else res['sentence'], normal_style),
                        res['source'],
                        f"{res['score']:.0f}%",
                        Paragraph(metadata_text, normal_style)
                    ])
                
                detail_table = Table(detail_data, colWidths=[3*inch, 1*inch, 0.6*inch, 1.6*inch])
                detail_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 9),
                    ('FONTSIZE', (0, 1), (-1, -1), 8),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ]))
                story.append(detail_table)
            else:
                story.append(Paragraph("<i>Tidak ada plagiarisme terdeteksi</i>", normal_style))
            
            # Legend
            story.append(Spacer(1, 0.3*inch))
            story.append(Paragraph("<b>Keterangan:</b>", normal_style))
            story.append(Paragraph("• Similaritas Global = Total kalimat terdeteksi plagiat / Total kalimat dokumen", normal_style))
            story.append(Paragraph("• Similaritas Lokal = Kalimat yang cocok dengan repository lokal / Total kalimat", normal_style))
            story.append(Paragraph("• Similaritas Internet = Kalimat yang cocok dengan internet / Total kalimat", normal_style))
            
            doc.build(story)
            print(f"✓ PDF report created: {output_path}")
            return output_path
            
        except Exception as e:
            print(f"✗ Error generating PDF: {e}")
            import traceback
            traceback.print_exc()
            return None